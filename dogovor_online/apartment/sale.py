from docx import Document
import os
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def string_to_data(context):
    print(context)

    doc = Document()
    style = doc.styles['Normal']
    # название шрифта
    style.font.name = 'Times New Roman'
    # размер шрифта
    style.font.size = Pt(14)
    # Звголовок договора
    header_1 = doc.add_paragraph()
    header_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = header_1.add_run(context['deal'] + ' ' + context['object'])
    run.font.size = Pt(22)
    # Данные договора
    items = (
    (context['all_params']['location'], context['all_params']['deal_number'], context['all_params']['date']),)
    # добавляем таблицу с одной строкой
    table = doc.add_table(cols=3, rows=1)
    for row in items:
        # добавляем строку с ячейками к объекту таблицы
        cells = table.add_row().cells
        # cells.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for i, item in enumerate(row):
            # вставляем данные в ячейки
            cells[i].text = str(item)
    doc.save('test.docx')
